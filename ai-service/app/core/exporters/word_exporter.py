import os
from typing import List, Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class WordExporter:
    def export(self, title: str, doc_type: str, steps: List[Dict[str, Any]], output_dir: str, filename: str) -> str:
        filepath = os.path.join(output_dir, f"{filename}.docx")
        doc = Document()

        # Title
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if doc_type:
            p = doc.add_paragraph()
            p.add_run(f"文档类型: ").bold = True
            p.add_run(doc_type)

        doc.add_paragraph()

        for step in steps:
            doc.add_heading(f"步骤 {step['step_no']}: {step['title']}", level=2)
            doc.add_paragraph(step['description'])
            doc.add_paragraph()

        doc.save(filepath)
        return filepath
